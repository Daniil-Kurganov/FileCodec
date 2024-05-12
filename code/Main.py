import sys
import os
import docx
import openpyxl
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from HammingCodec import hamming_codec_encode, hamming_codec_decode
from GUI import *

def  set_file_source_path() -> None:
    '''Установка пути файла обработки'''
    tuple_file_information = QFileDialog.getOpenFileName(caption = 'Выберите файл обработки',
                                                                   directory = 'C:/Users/User/PythonProjects/FileCodec/files',
                            filter = 'Текстовый документ (*.txt);; Документ Microsoft Word (*.docx);; Книга Microsoft Excel (*.xlsx)')
    ui.TextEditFileSourcePath.setText(tuple_file_information[0])
    return None
def set_directory_result_path() -> None:
    '''Установка пути каталога назначения'''
    directory = QFileDialog.getExistingDirectory(caption = 'Выберите каталог назначения',
                                                        directory = 'C:/Users/User/PythonProjects/FileCodec', options = QFileDialog.ShowDirsOnly)
    ui.TextEditFileResultPath.setText(directory)
    return None
def show_messages(float_message_key: float, string_message_text = None) -> None:
    '''Сообщений в зависимости от ключа. Формат словаря: {код: [иконка, название окна, заголовок, текст], ...}'''
    dictionary_messages_settings = {
        0.0: [QMessageBox.Information, 'Успех!', 'Кодирование завершилось успешно', 'Файл обработки успешно закодирован и сохранён.'],
        0.1: [QMessageBox.Information, 'Успех!', 'Декодирование завершилось успешно', 'Файл обработки успешно декодирован и сохранён.'],
        6.0: [QMessageBox.Warning, 'Внимание!', 'Файл назначения уже существует', 'Файл назначения будет перезаписан.'],
        13.0: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода путей файлов', 'Пути файлов совпадают.'],
        13.11: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода пути исходного файла', 'Файл обработки не существует.'],
        13.12: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода пути исходного файла', 'Расширение файла обработки не поддерживается.'],
        13.13: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода пути исходного файла', string_message_text],
        13.21: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода пути файла назначения', 'Директория назначения не существует.'],
        13.22: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода пути файла назначения', 'Расширение файла назначения не поддерживается.'],
        13.23: [QMessageBox.Critical, 'Ошибка!', 'Ошибка ввода пути файла назначения', string_message_text],
        13.3: [QMessageBox.Critical, 'Ошибка!', 'Ошибка работы кодека', string_message_text],
    }
    message = QMessageBox()
    message.setIcon(dictionary_messages_settings[float_message_key][0])
    message.setWindowTitle(dictionary_messages_settings[float_message_key][1])
    message.setText(dictionary_messages_settings[float_message_key][2])
    message.setInformativeText(dictionary_messages_settings[float_message_key][3])
    message.exec_()
    return None
def checking_file_paths(string_filepath_source: str, string_filepath_result: str) -> bool:
    '''Провекра корекстности введённых путей'''
    global string_extension
    if string_filepath_source == string_filepath_result:
        show_messages(13.0)
        return False
    try:
        if not os.path.isfile(string_filepath_source):
            show_messages(13.11)
            return False
        string_path, string_extension = os.path.splitext(string_filepath_source)
        if string_extension not in ['.txt', '.docx', '.xlsx']:
            show_messages(13.12)
            return False
    except Exception as error:
        show_messages(13.13, string_message_text = str(error))
        return False
    try:
        if not os.path.isdir(string_filepath_result[:string_filepath_result.rfind('/')]):
            show_messages(13.21)
            return False
        string_path, string_extension = os.path.splitext(string_filepath_result)
        if string_extension not in ['.txt', '.docx', '.xlsx']:
            show_messages(13.22)
            return False
        if os.path.isfile(string_filepath_result): show_messages(6.0)
    except Exception as error:
        show_messages(13.23, string_message_text = str(error))
        return False
    return True
def start_working_codec() -> None:
    '''Старт работы кодека'''
    global bool_operation_is_encode, string_extension
    string_filepath_source, string_filepath_result = ui.TextEditFileSourcePath.toPlainText(), ui.TextEditFileResultPath.toPlainText()
    if not checking_file_paths(string_filepath_source, string_filepath_result): return None
    int_counter_done_operations = 0
    try:
        if string_extension == '.txt':
            file_source = open(string_filepath_source, encoding='utf8')
            file_result = open(string_filepath_result, 'w', encoding='utf8')
            int_count_of_rows = sum(1 for string_current_row in file_source)
            file_source = open(string_filepath_source, encoding='utf8')
            for string_current_row in file_source:
                if bool_operation_is_encode: file_result.write(hamming_codec_encode(string_current_row[:-1]) + '\n')
                else: file_result.write(hamming_codec_decode(string_current_row[:-1]) + '\n')
                int_counter_done_operations += 1
                ui.ProgressBar.setProperty('value', ((int_counter_done_operations / int_count_of_rows) * 100))
            file_source.close()
            file_result.close()
        elif string_extension == '.docx':
            file_source, file_result = docx.Document(string_filepath_source), docx.Document()
            for paragraph_current in file_source.paragraphs:
                if bool_operation_is_encode: file_result.add_paragraph(hamming_codec_encode(paragraph_current.text))
                else: file_result.add_paragraph(hamming_codec_decode(paragraph_current.text))
                int_counter_done_operations += 1
                ui.ProgressBar.setProperty('value', (int_counter_done_operations / len(file_source.paragraphs)) * 100)
            file_result.save(string_filepath_result)
        elif string_extension == '.xlsx':
            file_source, file_result = openpyxl.load_workbook(string_filepath_source, read_only=True), openpyxl.Workbook()
            for sheet_current_source in file_source.worksheets:
                int_percentage_current_sheet_comlete = 100 / len(file_source.worksheets)
                sheet_current_result = file_result.create_sheet(title=sheet_current_source.title)
                for row_current_source in sheet_current_source.iter_rows():
                    int_percentage_current_row_comlete = int_percentage_current_sheet_comlete / sheet_current_source.max_row
                    for cell_current_source in row_current_source:
                        if cell_current_source.value is not None:
                            if bool_operation_is_encode:
                                sheet_current_result.cell(row=cell_current_source.row, column=cell_current_source.column,
                                                          value=hamming_codec_encode(str(cell_current_source.value)))
                            else:
                                sheet_current_result.cell(row=cell_current_source.row, column=cell_current_source.column,
                                                          value=hamming_codec_decode(str(cell_current_source.value)))
                        int_counter_done_operations += int_percentage_current_row_comlete / len(row_current_source)
                        ui.ProgressBar.setProperty('value', int_counter_done_operations)
            file_source.close()
            file_result.remove(file_result['Sheet'])
            file_result.save(string_filepath_result)
    except Exception as error:
        show_messages(13.3, string_message_text = str(error))
        return None
    if bool_operation_is_encode: show_messages(0.0)
    else: show_messages(0.1)
    return None
def set_encode_type_operation() -> None:
    '''Установка типа операции: кодирование'''
    global bool_operation_is_encode
    bool_operation_is_encode = True
    ui.PushButtonStart.setEnabled(True)
    ui.PushButtonStart.setText('Кодировать')
    return None
def set_decode_type_operation() -> None:
    '''Установка типа операции: декодирование'''
    global bool_operation_is_encode
    bool_operation_is_encode = False
    ui.PushButtonStart.setEnabled(True)
    ui.PushButtonStart.setText('Декодировать')
    return None

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    Window = QtWidgets.QMainWindow()
    ui = Ui_Window()
    ui.setupUi(Window)
    Window.show()
    ui.ToolButtonSetFileSourcePath.clicked.connect(set_file_source_path)
    ui.ToolButtonSetDirectoryResultPath.clicked.connect(set_directory_result_path)
    ui.RadioButtonEncodeType.clicked.connect(set_encode_type_operation)
    ui.RadioButtonDecodeType.clicked.connect(set_decode_type_operation)
    ui.PushButtonStart.clicked.connect(start_working_codec)
    sys.exit(app.exec_())


    # string_filepath_source = 'C:/Users/User/PythonProjects/FileCodec/files/ExcelFile.xlsx'
    # string_file_type = string_filepath_source[string_filepath_source.rfind('.') + 1:]
    # string_filepath_result = string_filepath_source[:string_filepath_source.rfind('/') + 1] + 'ExcelFileEncode' + '.' + string_file_type
    # bool_operation_is_encode = True
    # int_counter_done_operations = 0
    # if string_file_type == 'txt':
    #     file_source = open(string_filepath_source, encoding='utf8')
    #     file_result = open(string_filepath_result, 'w', encoding='utf8')
    #     int_count_of_rows = sum(1 for string_current_row in file_source)
    #     file_source = open(string_filepath_source, encoding='utf8')
    #     for string_current_row in file_source:
    #         if bool_operation_is_encode: file_result.write(hamming_codec_encode(string_current_row[:-1]) + '\n')
    #         else: file_result.write(hamming_codec_decode(string_current_row[:-1]) + '\n')
    #         int_counter_done_operations += 1
    #         print((int_counter_done_operations / int_count_of_rows) * 100)
    #     file_source.close()
    #     file_result.close()
    # elif string_file_type == 'docx':
    #     file_source, file_result = docx.Document(string_filepath_source), docx.Document()
    #     for paragraph_current in file_source.paragraphs:
    #         if bool_operation_is_encode: file_result.add_paragraph(hamming_codec_encode(paragraph_current.text))
    #         else: file_result.add_paragraph(hamming_codec_decode(paragraph_current.text))
    #         int_counter_done_operations += 1
    #         print((int_counter_done_operations / len(file_source.paragraphs)) * 100)
    #     file_result.save(string_filepath_result)
    # elif string_file_type == 'xlsx':
    #     file_source, file_result = openpyxl.load_workbook(string_filepath_source, read_only=True), openpyxl.Workbook()
    #     for sheet_current_source in file_source.worksheets:
    #         print(file_source.worksheets.index(sheet_current_source))
    #         int_percentage_current_sheet_comlete = 100 / len(file_source.worksheets)
    #         sheet_current_result = file_result.create_sheet(title=sheet_current_source.title)
    #         for row_current_source in sheet_current_source.iter_rows():
    #             int_percentage_current_row_comlete = int_percentage_current_sheet_comlete / sheet_current_source.max_row
    #             for cell_current_source in row_current_source:
    #                 if cell_current_source.value is not None:
    #                     if bool_operation_is_encode:
    #                         sheet_current_result.cell(row=cell_current_source.row, column=cell_current_source.column,
    #                                                   value=hamming_codec_encode(str(cell_current_source.value)))
    #                     else:
    #                         sheet_current_result.cell(row=cell_current_source.row, column=cell_current_source.column,
    #                                                   value=hamming_codec_decode(str(cell_current_source.value)))
    #                 int_counter_done_operations += int_percentage_current_row_comlete / len(row_current_source)
    #                 print(int_counter_done_operations)
    #     file_source.close()
    #     file_result.remove(file_result['Sheet'])
    #     file_result.save(string_filepath_result)
